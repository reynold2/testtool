# -*- coding: UTF-8 -*-
import docx
class Word_Docx():
    def __init__(self, path):
        self.doxcpath = path
        self.context = docx.Document(path)

    def replace_text(self, old_text, new_text):
        """替换段落文本内容"""
        self.__TraverseContext(self.context, old_text, new_text)

    def save(self, path=None):
        if path:
            self.context.save(path)
        else:
            self.context.save(self.doxcpath)

    def __delete__(self, instance):
        super(Word_Docx, self).__delete__()

    def replace_header(self,old_text, new_text):
        """替换页眉段落文本内容"""
        for section in self.context.sections:
            self.__TraverseContext(section.header, old_text, new_text)

    def replace_table(self,old_text, new_text):
        """替换表格段落文本内容"""
        for table in self.context.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.__TraverseContext(cell, old_text, new_text)

    def __TraverseContext(self, ContextObject, old_text, new_text):
        """替换段落基础方法"""
        for paragraph in ContextObject.paragraphs:
            if old_text in paragraph.text:
                inline = paragraph.runs
                for i in inline:
                    if old_text in i.text:
                        text = i.text.replace(old_text, new_text)
                        i.text = text

    def Replace_Mian(self, old_text, new_text):
        try:
            self.replace_text(old_text, new_text)
            self.replace_table(old_text, new_text)
            self.replace_header(old_text, new_text)
        except:
            print("文件替换字符出现异常")
        finally:
            self.save()

# if  __name__  == "__main__":
#     test = Word_Docx("E:\\新建 Microsoft Word 文档.docx")
#     test.Replace_Mian("大","小")
#     test.save()

import os
import win32com
from win32com.client import Dispatch
# 处理Word文档的类
class RemoteWord:
    def __init__(self, filename=None):
        self.xlApp = win32com.client.Dispatch('Word.Application')  # 此处使用的是Dispatch，原文中使用的DispatchEx会报错
        self.xlApp.Visible = 0  # 后台运行，不显示
        self.xlApp.DisplayAlerts = 0 # 不警告
        if filename:
            self.filename = filename
            if os.path.exists(self.filename):
                self.doc = self.xlApp.Documents.Open(filename)
            else:
                self.doc = self.xlApp.Documents.Add() # 创建新的文档
                self.doc.SaveAs(filename)
        else:
            self.doc = self.xlApp.Documents.Add()
            self.filename = ''
    def add_doc_end(self, string):
        '''在文档末尾添加内容'''
        rangee = self.doc.Range()
        rangee.InsertAfter('\n' + string)
    def add_doc_start(self, string):
        """在文档开头添加内容"""
        rangee = self.doc.Range(0, 0)
        rangee.InsertBefore(string + '\n')
    def insert_doc(self, insertPos, string):
        '''在文档insertPos位置添加内容'''

        rangee = self.doc.Range(0, insertPos)
        if (insertPos == 0):
            rangee.InsertAfter(string)
        else:
            rangee.InsertAfter('\n' + string)
    def replace_doc(self, string, new_string):
        '''替换文字'''
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, False, False, False, True, 1, True, new_string, 2)
        # (string--搜索文本,
        # True--区分大小写,
        # True--完全匹配的单词，并非单词中的部分（全字匹配）,
        # True--使用通配符,
        # True--同音,
        # True--查找单词的各种形式,
        # True--向文档尾部搜索,
        # 1,
        # True--带格式的文本,
        # new_string--替换文本,
        # 2--替换个数（全部替换）
    def replace_docs(self, string, new_string):
        '''采用通配符匹配替换'''
        self.xlApp.Selection.Find.ClearFormatting()
        self.xlApp.Selection.Find.Replacement.ClearFormatting()
        self.xlApp.Selection.Find.Execute(string, False, False, True, False, False, False, 1, False, new_string, 2)
    def header(self,string, new_string):
        Sections=self.xlApp.ActiveDocument.Sections
        for Section in Sections:
            for Header in Section.Headers:
                Header.Range.Find.ClearFormatting()
                Header.Range.Find.Replacement.ClearFormatting()
                Header.Range.Find.Execute(string, False, False, False, False, False, True, 1,
                                                          False, new_string, 2)
    def save(self):
        '''保存文档'''
        self.doc.Save()

    def save_as(self, filename):
        '''文档另存为'''
        self.doc.SaveAs(filename)

    def close(self):
        '''保存文件、关闭文件'''
        self.save()
        self.xlApp.Documents.Close()
        self.xlApp.Quit()

    def word_replace(self, string, new_string):
        '''替换内容和页眉'''
        self.header(string, new_string)
        self.replace_doc(string, new_string)

# if __name__ == '__main__':
#     path = 'F:\\testtool[dev]\\cproject\\FileTool\pyword\\处级QBCL系统改造个性软件配置项测试说明.docx'
#     doc = RemoteWord(path)
#     doc.word_replace('测试', '淼')
#     doc.close()

#-*- encoding: utf8 -*-
import win32com.client
import os
import time

class RemoteExcel():
    """对excel表格的操作

    """
    def __init__(self, filename=None):
        """初始化函数

        Args:
            filename: 要进行操作的文件名，如果存在此文件则打开，不存在则新建
                        此文件名一般包含路径

        """
        self.xlApp=win32com.client.Dispatch('Excel.Application')
        self.xlApp.Visible=0
        self.xlApp.DisplayAlerts=0    #后台运行，不显示，不警告
        if filename:
            self.filename=filename
            if os.path.exists(self.filename):
                self.xlBook=self.xlApp.Workbooks.Open(filename)
            else:
                self.xlBook = self.xlApp.Workbooks.Add()    #创建新的Excel文件
                self.xlBook.SaveAs(self.filename)
        else:
            self.xlBook=self.xlApp.Workbooks.Add()
            self.filename=''

    def get_cell(self, row, col, sheet=None):
        """读取单元格的内容

        Args:
            row: 行
            col: 列
            sheet: 表格名（不是文件名）

        """
        if sheet:
            sht = self.xlBook.Worksheets(sheet)
        else:
            sht = self.xlApp.ActiveSheet
        return sht.Cells(row, col).Value

    def set_cell(self, sheet, row, col, value):
        """向表格单元格写入

        Args:
            sheet: 表格名（不是文件名）
            row: 行
            col: 列
            value: 定入内容
        """
        try:
            sht = self.xlBook.Worksheets(sheet)
        except:
            self.new_sheet(sheet)
            sht = self.xlBook.Worksheets(sheet)
        sht.Cells(row, col).Value = value

    def save(self, newfilename=None):
        """保存表格"""
        if newfilename:
            self.filename = newfilename
            self.xlBook.SaveAs(self.filename)
        else:
            self.xlBook.Save()

    def close(self):
        """保存表格、关闭表格，结束操作"""
        self.save()
        self.xlBook.Close(SaveChanges=0)
        del self.xlApp

    def new_sheet(self, newSheetName):
        """新建一个新表格"""
        sheet = self.xlBook.Worksheets.Add()
        sheet.Name = newSheetName
        sheet.Activate()

    def active_sheet(self):
        return self.xlApp.ActiveSheet

    def re_Excel(self,old,new):
        for Worksheet in self.xlBook.Worksheets:
            Rows=Worksheet.UsedRange.Rows.Count
            Columns=Worksheet.UsedRange.Columns.Count
            for Row in range(Rows):
                for Column in range(Columns):
                    print(Row+1,Column+1)
                    text=Worksheet.Cells(Row+1, Column+1).Value
                    if type(text) is float:
                        Worksheet.Cells(Row+1, Column+1).Value=str(int(text)).replace(str(old),str(new))
                    elif type(text) is str:
                        if old in text:
                            Worksheet.Cells(Row+1, Column+1).Value=text.replace(old,new)
                    else:
                        if str(old) in str(text):
                            Worksheet.Cells(Row+1, Column+1).Value=str(text).replace(old,new)
        self.save()


if __name__=='__main__':
    #example1
    TODAY = time.strftime('%Y-%m-%d',    time.localtime(time.time()))
    excel = RemoteExcel('E:\\2019内审检查表(8章)_r.xlsx')

    excel.re_Excel(".",")")
    # excel.set_cell('old',1,1, "111")
    # print (excel.get_cell(1,1,'old'))
    excel.close()
